"""
Document ingestion pipeline:
  File → text extraction → chunking → embedding → Qdrant indexing

Supported formats:
  .pdf   — PyPDFLoader
  .docx  — python-docx (paragraphes + tableaux), fallback Docx2txtLoader
  .txt   — TextLoader (utf-8)
  .md    — TextLoader (utf-8)  [UnstructuredMarkdownLoader is too heavy]

Also maintains a per-user BM25 corpus persisted to /data/bm25/<user_id>.pkl.
"""
from __future__ import annotations

import hashlib
import logging
import os
import pickle
import re
from pathlib import Path
from typing import Any

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams

from .config import (
    BM25_DIR,
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    CHUNKER,
    EMBEDDING_DIM,
    EMBEDDING_MODEL,
    QDRANT_API_KEY,
    QDRANT_COLLECTION,
    QDRANT_URL,
    bm25_file,
)
from .semantic_chunker import (
    CHUNKER_VERSION,
    semantic_chunk_documents,
)

logger = logging.getLogger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".xlsx", ".xls"}

# ---------------------------------------------------------------------------
# Shared singletons (lazy initialisation)
# ---------------------------------------------------------------------------

_embeddings: HuggingFaceEmbeddings | None = None
_qdrant_client: QdrantClient | None = None

# Per-user in-memory BM25 corpora:  {user_id: list[dict]}
_bm25_corpora: dict[str, list[dict[str, Any]]] = {}


# ---------------------------------------------------------------------------
# Collection name helpers
# ---------------------------------------------------------------------------


def sanitize_collection_name(user_id: str) -> str:
    """
    Return a safe Qdrant collection name for the given user_id.

    Rules: lowercase, replace non-alphanumeric with '_', prefix 'rag_', max 40 chars.
    """
    safe = re.sub(r"[^a-zA-Z0-9]", "_", user_id).lower()
    name = f"rag_{safe}"
    return name[:40]


def _collection_for_user(user_id: str) -> str:
    """Convenience alias used in this module."""
    return sanitize_collection_name(user_id)


# ---------------------------------------------------------------------------
# Singleton accessors
# ---------------------------------------------------------------------------


EMBED_BATCH_SIZE: int = int(os.getenv("EMBED_BATCH_SIZE", "8"))
EMBED_MAX_SEQ_LENGTH: int = int(os.getenv("EMBED_MAX_SEQ_LENGTH", "4096"))


def get_embeddings() -> HuggingFaceEmbeddings:
    """Singleton bge-m3 hardé pour usage CPU.

    - batch_size petit (défaut 8) pour éviter les pics mémoire CPU.
    - max_seq_length limité à 4096 tokens (au lieu des 8192 max bge-m3) pour
      éviter qu'un chunk pathologique (HTML mal nettoyé, page d'erreur, longue
      chaîne sans espace) ne fasse exploser le tokenizer ou la passe forward.
    - show_progress_bar et convert_to_numpy ne sont PAS passés via encode_kwargs
      car langchain_huggingface les passe déjà en interne à
      SentenceTransformer.encode() — un doublon provoque
      "got multiple values for keyword argument".
    """
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            encode_kwargs={
                "normalize_embeddings": True,
                "batch_size": EMBED_BATCH_SIZE,
            },
        )
        # SentenceTransformer expose max_seq_length sur l'instance interne.
        # On le réduit à 4096 pour borner le coût de l'embedding par chunk.
        try:
            _embeddings._client.max_seq_length = EMBED_MAX_SEQ_LENGTH
            logger.info(
                "Embeddings bge-m3 prêts (batch_size=%d, max_seq_length=%d)",
                EMBED_BATCH_SIZE,
                EMBED_MAX_SEQ_LENGTH,
            )
        except Exception as exc:  # pragma: no cover — défensif
            logger.warning(
                "Impossible de fixer max_seq_length sur le SentenceTransformer: %s",
                exc,
            )
    return _embeddings


def get_qdrant_client(qdrant_url: str = QDRANT_URL) -> QdrantClient:
    global _qdrant_client
    if _qdrant_client is None:
        if QDRANT_API_KEY:
            _qdrant_client = QdrantClient(url=qdrant_url, api_key=QDRANT_API_KEY)
        else:
            _qdrant_client = QdrantClient(url=qdrant_url)
    return _qdrant_client


# ---------------------------------------------------------------------------
# Collection management
# ---------------------------------------------------------------------------


def ensure_collection(client: QdrantClient, collection_name: str = QDRANT_COLLECTION) -> None:
    """Create the Qdrant collection if it does not exist yet."""
    existing = [c.name for c in client.get_collections().collections]
    if collection_name not in existing:
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(
                size=EMBEDDING_DIM,
                distance=Distance.COSINE,
            ),
        )
        logger.info("Collection '%s' created.", collection_name)


# ---------------------------------------------------------------------------
# Per-user BM25 corpus management
# ---------------------------------------------------------------------------


def _ensure_bm25_dir() -> None:
    """Create the BM25 directory if it doesn't exist."""
    Path(BM25_DIR).mkdir(parents=True, exist_ok=True)


def load_bm25_corpus(user_id: str = "default") -> list[dict[str, Any]]:
    """
    Load and return the BM25 corpus for the given user.
    Loads from disk on first call; subsequent calls use in-memory cache.
    """
    if user_id in _bm25_corpora:
        return _bm25_corpora[user_id]

    corpus: list[dict[str, Any]] = []
    pkl_path = bm25_file(user_id)
    if Path(pkl_path).exists():
        try:
            with open(pkl_path, "rb") as fh:
                corpus = pickle.load(fh)
            logger.info(
                "Loaded BM25 corpus for user '%s' with %d chunks.", user_id, len(corpus)
            )
        except Exception as exc:
            logger.warning("Could not load BM25 corpus for user '%s': %s", user_id, exc)
            corpus = []

    _bm25_corpora[user_id] = corpus
    return corpus


def save_bm25_corpus(user_id: str = "default") -> None:
    """Persist the in-memory BM25 corpus for user_id to disk."""
    _ensure_bm25_dir()
    corpus = _bm25_corpora.get(user_id, [])
    pkl_path = bm25_file(user_id)
    try:
        with open(pkl_path, "wb") as fh:
            pickle.dump(corpus, fh)
    except Exception as exc:
        logger.warning("Could not save BM25 corpus for user '%s': %s", user_id, exc)


def reset_bm25_corpus(user_id: str = "default") -> None:
    """Clear in-memory and on-disk BM25 corpus for a user."""
    _bm25_corpora[user_id] = []
    pkl_path = bm25_file(user_id)
    if Path(pkl_path).exists():
        os.remove(pkl_path)


# ---------------------------------------------------------------------------
# Loader selection by extension
# ---------------------------------------------------------------------------


def _load_excel_document(file_path: str, ext: str, source_name: str):
    """Extrait le texte d'un classeur Excel (xlsx/xls) en concaténant les lignes.

    Pour chaque feuille, émet un en-tête `## Feuille : {nom}` puis chaque ligne
    sous forme `cell1 | cell2 | ...` (cellules vides ignorées). Retourne un
    Document unique : le chunker sémantique le découpera ensuite par taille.
    """
    from langchain_core.documents import Document

    parts: list[str] = []
    if ext == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True, read_only=True)
        try:
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"## Feuille : {sheet_name}")
                for row in ws.iter_rows(values_only=True):
                    cells = [
                        str(c).strip()
                        for c in row
                        if c is not None and str(c).strip()
                    ]
                    if cells:
                        parts.append(" | ".join(cells))
        finally:
            wb.close()
    elif ext == ".xls":
        import xlrd

        wb = xlrd.open_workbook(file_path)
        for sheet in wb.sheets():
            parts.append(f"## Feuille : {sheet.name}")
            for row_idx in range(sheet.nrows):
                row = sheet.row_values(row_idx)
                cells = [
                    str(c).strip()
                    for c in row
                    if c is not None and str(c).strip() != ""
                ]
                if cells:
                    parts.append(" | ".join(cells))
    else:
        raise ValueError(f"Extension Excel non gérée : {ext}")

    text = "\n".join(parts)
    return [Document(page_content=text, metadata={"source": source_name, "page": 1})]


def _load_docx_document(file_path: str, source_name: str):
    """Extrait le texte ET les tableaux d'un .docx en préservant l'ordre.

    Utilise python-docx (au lieu de docx2txt) qui ne lit que le texte plat
    et perd les tableaux. Pour des dossiers de conception SIRH qui sont
    majoritairement structurés en tableaux, le gain est important :
    typiquement 3-5× plus de chunks utiles.

    Format de sortie (markdown-like) — le chunker sémantique v2 reconnaît
    les # comme breaks de section :
        # Heading 1
        ## Heading 2
        Paragraphe...

        | col1 | col2 |
        | val1 | val2 |
        | ...  | ...  |

    En cas d'erreur python-docx, fallback sur Docx2txtLoader.
    """
    from langchain_core.documents import Document

    try:
        import docx as pydocx  # python-docx
        from docx.oxml.ns import qn  # type: ignore
    except ImportError:
        # Fallback : si python-docx pas installé, on retombe sur l'ancien loader
        from langchain_community.document_loaders import Docx2txtLoader
        return Docx2txtLoader(file_path).load()

    try:
        doc = pydocx.Document(file_path)
    except Exception as exc:
        logger.warning(
            "[docx] python-docx a échoué sur '%s' : %s. Fallback docx2txt.",
            source_name, exc,
        )
        from langchain_community.document_loaders import Docx2txtLoader
        return Docx2txtLoader(file_path).load()

    # Mappe chaque <w:p> et <w:tbl> du body dans l'ordre du document.
    # python-docx fournit doc.paragraphs et doc.tables séparément ;
    # on parcourt l'élément XML pour préserver l'ordre.
    parts: list[str] = []
    P_TAG = qn("w:p")
    TBL_TAG = qn("w:tbl")
    paragraphs_by_id = {p._element: p for p in doc.paragraphs}
    tables_by_id = {t._element: t for t in doc.tables}

    def _md_heading_level(style_name: str) -> int:
        """Retourne 1, 2, 3… si le style commence par 'Heading' / 'Titre',
        sinon 0 (= paragraphe normal)."""
        if not style_name:
            return 0
        s = style_name.strip().lower()
        for prefix in ("heading ", "titre "):
            if s.startswith(prefix):
                tail = s[len(prefix):].strip()
                try:
                    n = int(tail)
                    return max(1, min(6, n))
                except ValueError:
                    return 1
        return 0

    for child in doc.element.body.iterchildren():
        if child.tag == P_TAG:
            p = paragraphs_by_id.get(child)
            if p is None:
                continue
            text = (p.text or "").strip()
            if not text:
                continue
            level = _md_heading_level(p.style.name if p.style else "")
            if level > 0:
                parts.append("\n" + ("#" * level) + " " + text + "\n")
            else:
                parts.append(text)
        elif child.tag == TBL_TAG:
            tbl = tables_by_id.get(child)
            if tbl is None:
                continue
            rows_text: list[str] = []
            for row in tbl.rows:
                cells = [
                    " ".join((c.text or "").split()).strip() for c in row.cells
                ]
                # Ignore les lignes 100% vides
                if any(cells):
                    rows_text.append("| " + " | ".join(cells) + " |")
            if rows_text:
                parts.append("")  # ligne vide avant le tableau
                parts.extend(rows_text)
                parts.append("")  # ligne vide après

    full_text = "\n".join(parts).strip()
    if not full_text:
        # Document vide ou contenu non extractible → fallback docx2txt au cas où
        logger.warning(
            "[docx] python-docx n'a rien extrait de '%s', fallback docx2txt.",
            source_name,
        )
        from langchain_community.document_loaders import Docx2txtLoader
        return Docx2txtLoader(file_path).load()

    return [Document(page_content=full_text, metadata={"source": source_name, "page": 1})]


def _load_documents(file_path: str, ext: str, source_name: str | None = None):
    """
    Load a document using the appropriate loader for the file extension.
    Returns a list of LangChain Document objects.
    """
    ext = ext.lower()
    if ext == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(file_path)
    elif ext == ".docx":
        # Loader custom python-docx (préserve paragraphes + tableaux). Fallback
        # docx2txt intégré dans le helper si python-docx n'est pas dispo ou
        # n'arrive pas à parser.
        return _load_docx_document(file_path, source_name or file_path)
    elif ext in {".txt", ".md"}:
        from langchain_community.document_loaders import TextLoader
        loader = TextLoader(file_path, encoding="utf-8")
    elif ext in {".xlsx", ".xls"}:
        return _load_excel_document(file_path, ext, source_name or file_path)
    else:
        raise ValueError(
            f"Format non supporté : '{ext}'. "
            f"Formats acceptés : {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return loader.load()


# ---------------------------------------------------------------------------
# Main ingestion function (per-user)
# ---------------------------------------------------------------------------


def ingest_file(
    file_path: str,
    source_name: str,
    user_id: str = "default",
    qdrant_url: str = QDRANT_URL,
) -> int:
    """
    Ingest a document file into the user's Qdrant collection (dense) and
    their BM25 corpus (sparse).

    Supported formats: PDF, DOCX, TXT, MD, XLSX, XLS.

    Returns the number of chunks indexed.
    """
    collection_name = _collection_for_user(user_id)

    # Ensure corpus is loaded
    load_bm25_corpus(user_id)

    # Determine extension
    ext = Path(source_name).suffix.lower()
    if not ext:
        ext = Path(file_path).suffix.lower()

    # 1. Load document
    pages = _load_documents(file_path, ext, source_name)
    logger.info("Loaded %d page(s)/section(s) from '%s'.", len(pages), source_name)

    # 2. Split into chunks — semantic (v3.9.0 default) or legacy size-based
    embeddings = get_embeddings()
    if CHUNKER == "semantic":
        logger.info(
            "Chunking '%s' with semantic+structure-aware chunker (%s).",
            source_name, CHUNKER_VERSION,
        )
        docs = semantic_chunk_documents(pages, embeddings.embed_documents)
    else:
        logger.info("Chunking '%s' with legacy RecursiveCharacterTextSplitter.", source_name)
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        docs = splitter.split_documents(pages)

    # 3. Enrich metadata
    doc_hash = hashlib.md5(source_name.encode()).hexdigest()[:8]
    for i, doc in enumerate(docs):
        doc.metadata["source"] = source_name
        doc.metadata["chunk_id"] = f"{doc_hash}_{i}"
        # PyPDFLoader sets 'page' (0-indexed) — convert to 1-indexed
        if "page" in doc.metadata:
            doc.metadata["page"] = int(doc.metadata["page"]) + 1
        else:
            doc.metadata["page"] = 1
        # Ensure chunker_version is stamped (semantic chunker already sets it)
        doc.metadata.setdefault("chunker_version", CHUNKER_VERSION if CHUNKER == "semantic" else "legacy")

    # 4. Embed & store in user's Qdrant collection
    client = get_qdrant_client(qdrant_url)
    ensure_collection(client, collection_name)

    vector_store = QdrantVectorStore(
        client=client,
        collection_name=collection_name,
        embedding=embeddings,
    )
    vector_store.add_documents(docs)

    # 5. Store in user's BM25 corpus
    corpus = _bm25_corpora.setdefault(user_id, [])
    for doc in docs:
        corpus.append(
            {
                "id": doc.metadata["chunk_id"],
                "text": doc.page_content,
                "metadata": doc.metadata,
            }
        )

    save_bm25_corpus(user_id)
    logger.info(
        "Indexed %d chunks from '%s' for user '%s' (collection '%s'). BM25 size: %d.",
        len(docs),
        source_name,
        user_id,
        collection_name,
        len(corpus),
    )
    return len(docs)


# ---------------------------------------------------------------------------
# Helper functions for user collections
# ---------------------------------------------------------------------------


def get_all_collections(qdrant_url: str = QDRANT_URL) -> dict[str, int]:
    """Return a dict of {collection_name: vector_count} for all collections."""
    try:
        client = get_qdrant_client(qdrant_url)
        collections = client.get_collections().collections
        result = {}
        for col in collections:
            try:
                info = client.get_collection(col.name)
                result[col.name] = info.points_count or 0
            except Exception:
                result[col.name] = 0
        return result
    except Exception:
        return {}


def list_user_documents(user_id: str = "default", qdrant_url: str = QDRANT_URL) -> list[dict[str, Any]]:
    """Return a list of unique documents indexed for user_id.

    Each item: {"source": str, "chunks": int}. Sourced from the BM25 corpus
    (persisted to disk), which mirrors what's in Qdrant but is faster to read.
    Falls back to Qdrant scroll if BM25 is empty but the collection has points.
    """
    corpus = load_bm25_corpus(user_id)
    counts: dict[str, int] = {}
    for entry in corpus:
        src = entry.get("metadata", {}).get("source") or "unknown"
        counts[src] = counts.get(src, 0) + 1

    # Fallback: if BM25 is empty but Qdrant has data (edge case), scroll Qdrant
    if not counts:
        try:
            collection_name = _collection_for_user(user_id)
            client = get_qdrant_client(qdrant_url)
            existing = [c.name for c in client.get_collections().collections]
            if collection_name in existing:
                offset = None
                while True:
                    points, offset = client.scroll(
                        collection_name=collection_name,
                        limit=256,
                        with_payload=True,
                        with_vectors=False,
                        offset=offset,
                    )
                    for p in points:
                        payload = p.payload or {}
                        meta = payload.get("metadata", {}) or {}
                        src = meta.get("source") or payload.get("source") or "unknown"
                        counts[src] = counts.get(src, 0) + 1
                    if offset is None:
                        break
        except Exception as exc:
            logger.warning("Could not scroll Qdrant for user '%s': %s", user_id, exc)

    return [{"source": s, "chunks": c} for s, c in sorted(counts.items())]


def delete_document_by_source(
    source: str,
    user_id: str = "default",
    qdrant_url: str = QDRANT_URL,
) -> dict:
    """Delete all chunks of a given source file for a user.

    Removes matching points from Qdrant (filter on metadata.source) and
    filters the BM25 corpus. Returns counts of removed items.
    """
    from qdrant_client.models import Filter, FieldCondition, MatchValue, FilterSelector

    collection_name = _collection_for_user(user_id)
    client = get_qdrant_client(qdrant_url)
    existing = [c.name for c in client.get_collections().collections]

    qdrant_deleted = 0
    if collection_name in existing:
        # Count first (for reporting), then delete by filter on metadata.source
        qdrant_filter = Filter(
            must=[
                FieldCondition(
                    key="metadata.source",
                    match=MatchValue(value=source),
                )
            ]
        )
        try:
            count_resp = client.count(
                collection_name=collection_name,
                count_filter=qdrant_filter,
                exact=True,
            )
            qdrant_deleted = int(count_resp.count)
        except Exception as exc:
            logger.warning("Could not count points for source '%s': %s", source, exc)

        try:
            client.delete(
                collection_name=collection_name,
                points_selector=FilterSelector(filter=qdrant_filter),
            )
        except Exception as exc:
            logger.error("Qdrant delete failed for source '%s': %s", source, exc)
            raise

    # Update BM25 corpus in-memory + on-disk
    corpus = load_bm25_corpus(user_id)
    before = len(corpus)
    filtered = [
        entry
        for entry in corpus
        if (entry.get("metadata", {}) or {}).get("source") != source
    ]
    bm25_deleted = before - len(filtered)
    _bm25_corpora[user_id] = filtered
    save_bm25_corpus(user_id)

    logger.info(
        "Deleted source '%s' for user '%s' (qdrant=%d, bm25=%d).",
        source,
        user_id,
        qdrant_deleted,
        bm25_deleted,
    )
    return {
        "source": source,
        "qdrant_deleted": qdrant_deleted,
        "bm25_deleted": bm25_deleted,
    }


def reset_collection(qdrant_url: str = QDRANT_URL, user_id: str = "default") -> None:
    """Delete and recreate the user's Qdrant collection, and reset their BM25 corpus."""
    collection_name = _collection_for_user(user_id)
    client = get_qdrant_client(qdrant_url)
    existing = [c.name for c in client.get_collections().collections]
    if collection_name in existing:
        client.delete_collection(collection_name)
    ensure_collection(client, collection_name)
    reset_bm25_corpus(user_id)
    logger.info("Collection '%s' reset for user '%s'.", collection_name, user_id)
